"""
    Local-first media indexing and deduplication CLI.
    Copyright (C) 2026  Dario Palladino

    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU General Public License as published by
    the Free Software Foundation, either version 3 of the License, or
    (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU General Public License for more details.

    You should have received a copy of the GNU General Public License
    along with this program.  If not, see <http://www.gnu.org/licenses/>.
"""
"""Document metadata extractor for PDF, DOCX, TXT, Markdown."""
from __future__ import annotations

from pathlib import Path
from typing import Any

import structlog

from mediactl.metadata.base import MetadataPlugin

log = structlog.get_logger()

DOC_EXTS = {"pdf", "docx", "doc", "txt", "md", "markdown", "rtf", "odt"}


class DocumentMetadataPlugin(MetadataPlugin):
    """Extract title, author, page count, text preview from documents."""

    def supports(self, file_type: str) -> bool:
        return file_type.lower().lstrip(".") in DOC_EXTS

    def process(self, file_path: Path) -> dict[str, Any]:
        ext = file_path.suffix.lower().lstrip(".")
        if ext == "pdf":
            return self._process_pdf(file_path)
        elif ext in ("docx", "doc"):
            return self._process_docx(file_path)
        elif ext in ("txt", "md", "markdown"):
            return self._process_text(file_path)
        return {"error": f"Unsupported document type: {ext}"}

    def _process_pdf(self, file_path: Path) -> dict[str, Any]:
        result: dict[str, Any] = {}
        try:
            import fitz  # PyMuPDF  # type: ignore[import-untyped]

            doc = fitz.open(str(file_path))
            meta = doc.metadata
            result["page_count"] = doc.page_count
            result["title"] = meta.get("title", "")
            result["author"] = meta.get("author", "")
            result["subject"] = meta.get("subject", "")
            result["creator"] = meta.get("creator", "")
            result["creation_date"] = meta.get("creationDate", "")
            result["modification_date"] = meta.get("modDate", "")

            # Text preview from first page
            if doc.page_count > 0:
                first_page = doc[0]
                result["text_preview"] = first_page.get_text()[:500]

            doc.close()
        except ImportError:
            result["error"] = "pymupdf not installed"
        except Exception as exc:
            log.warning("metadata.doc.pdf_error", path=str(file_path), error=str(exc))
            result["error"] = str(exc)
        return result

    def _process_docx(self, file_path: Path) -> dict[str, Any]:
        result: dict[str, Any] = {}
        try:
            import docx  # python-docx  # type: ignore[import-untyped]

            doc = docx.Document(str(file_path))
            core = doc.core_properties
            result["title"] = core.title or ""
            result["author"] = core.author or ""
            result["subject"] = core.subject or ""
            result["created"] = str(core.created) if core.created else ""
            result["modified"] = str(core.modified) if core.modified else ""
            result["paragraph_count"] = len(doc.paragraphs)

            # Text preview
            preview = " ".join(p.text for p in doc.paragraphs[:5] if p.text)
            result["text_preview"] = preview[:500]

        except ImportError:
            result["error"] = "python-docx not installed"
        except Exception as exc:
            log.warning("metadata.doc.docx_error", path=str(file_path), error=str(exc))
            result["error"] = str(exc)
        return result

    def _process_text(self, file_path: Path) -> dict[str, Any]:
        result: dict[str, Any] = {}
        try:
            content = file_path.read_text(encoding="utf-8", errors="replace")
            lines = content.splitlines()
            result["line_count"] = len(lines)
            result["char_count"] = len(content)
            result["text_preview"] = content[:500]

            # For markdown: try to extract title from first heading
            ext = file_path.suffix.lower().lstrip(".")
            if ext in ("md", "markdown"):
                for line in lines:
                    if line.startswith("# "):
                        result["title"] = line.lstrip("# ").strip()
                        break

        except Exception as exc:
            log.warning("metadata.doc.text_error", path=str(file_path), error=str(exc))
            result["error"] = str(exc)
        return result
